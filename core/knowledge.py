"""本地知识库 —— 零依赖的轻量 RAG（FTS5 全文检索）。

设计取舍：
- 不引入向量库 / embedding 模型，纯靠 SQLite FTS5 的 LIKE 子串匹配 + 命中次数排序。
  原因：用户机器是 16GB + 无 CUDA，embedding 模型会额外吃资源；而个人知识库通常规模不大，
  子串检索对中文（尤其是「按关键词找资料」这种用法）足够好用，且完全离线、零额外依赖。
- 文档切块后存进 FTS5 表，提问时对「最后一条用户消息」做检索，命中块拼进 system 上下文。

支持格式：
- 纯文本类（直接读）：txt / md / py / js / ts / json / csv / log / yaml / yml /
  html / css / java / go / rs / c / cpp / h / sh / bat / ps1 / sql
- docx / pdf：若本机装了对应解析库则自动用，否则返回友好提示（不强制加依赖）。
"""

from __future__ import annotations

import os
import re
import sqlite3
import time
from typing import Any

from .store import data_dir

CHUNK_SIZE = 700          # 每块约 700 字
CHUNK_OVERLAP = 100       # 块间重叠，避免一句话被切断
TOP_K = 4                 # 每次最多注入 4 块

_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".json", ".csv", ".log", ".yaml", ".yml", ".html", ".htm", ".css",
    ".java", ".go", ".rs", ".c", ".cpp", ".cc", ".h", ".hpp",
    ".sh", ".bash", ".bat", ".ps1", ".sql", ".toml", ".ini", ".cfg",
}


def _split(text: str) -> list[str]:
    """按段落切块，超长段落再硬切，块间留重叠。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    paras = text.split("\n")
    chunks: list[str] = []
    buf = ""
    for p in paras:
        if len(buf) + len(p) + 1 <= CHUNK_SIZE:
            buf = (buf + "\n" + p).strip() if buf else p.strip()
        else:
            if buf:
                chunks.append(buf)
            if len(p) > CHUNK_SIZE:
                step = max(1, CHUNK_SIZE - CHUNK_OVERLAP)
                for i in range(0, len(p), step):
                    seg = p[i : i + CHUNK_SIZE].strip()
                    if seg:
                        chunks.append(seg)
                buf = ""
            else:
                buf = p.strip()
    if buf:
        chunks.append(buf)
    return [c for c in chunks if c.strip()]


class Knowledge:
    def __init__(self, db_path: str | None = None):
        if db_path is None:
            db_path = os.path.join(data_dir(), "knowledge.db")
        self.db_path = db_path
        self._init()

    # ---------------------------------------------------------------- 内部
    def _conn(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path, timeout=15)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA journal_mode=WAL")
        return conn

    def _init(self) -> None:
        with self._conn() as c:
            c.execute(
                """
                CREATE TABLE IF NOT EXISTS docs (
                    id         INTEGER PRIMARY KEY AUTOINCREMENT,
                    name       TEXT NOT NULL,
                    path       TEXT NOT NULL DEFAULT '',
                    kind       TEXT NOT NULL DEFAULT 'text',
                    size       INTEGER NOT NULL DEFAULT 0,
                    created_at INTEGER NOT NULL
                )
                """
            )
            c.execute(
                """
                CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts
                USING fts5(content, doc_id UNINDEXED, idx UNINDEXED)
                """
            )

    # ---------------------------------------------------------------- 导入
    def import_file(self, path: str) -> dict[str, Any]:
        """导入一个文件。返回 {ok, name, chunks, error}。"""
        if not os.path.isfile(path):
            return {"ok": False, "error": f"文件不存在：{path}"}
        ext = os.path.splitext(path)[1].lower()
        name = os.path.basename(path)
        try:
            if ext in _TEXT_EXTS:
                text = self._read_text(path)
                kind = "text"
            elif ext == ".docx":
                text = self._read_docx(path)
                kind = "docx"
            elif ext == ".pdf":
                text = self._read_pdf(path)
                kind = "pdf"
            elif ext in (".xlsx", ".xls"):
                text = self._read_xlsx(path)
                kind = "sheet"
            else:
                return {
                    "ok": False,
                    "error": f"暂不支持的格式：{ext}（支持 txt/md/代码/docx/pdf）",
                }
        except ImportError as exc:
            return {
                "ok": False,
                "error": f"解析 {ext} 需要额外库：{exc}。请先 pip install 对应库。",
            }
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "error": f"读取失败：{exc}"}

        text = (text or "").strip()
        if not text:
            return {"ok": False, "error": "文件内容为空或无法提取文本。"}

        chunks = _split(text)
        with self._conn() as c:
            cur = c.execute(
                "INSERT INTO docs(name,path,kind,size,created_at) VALUES (?,?,?,?,?)",
                (name, path, kind, len(text), int(time.time() * 1000)),
            )
            doc_id = cur.lastrowid
            c.executemany(
                "INSERT INTO chunks_fts(content, doc_id, idx) VALUES (?,?,?)",
                [(ch, doc_id, i) for i, ch in enumerate(chunks)],
            )
        return {"ok": True, "name": name, "chunks": len(chunks), "doc_id": doc_id}

    @staticmethod
    def _read_text(path: str) -> str:
        # 依次尝试常见中文编码，避免乱码
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        # 兜底：忽略错误的字节
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def _read_docx(path: str) -> str:
        from docx import Document  # type: ignore

        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text for c in row.cells if c.text.strip()))
        return "\n\n".join(parts)

    @staticmethod
    def _read_pdf(path: str) -> str:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError:
            from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(path)
        return "\n\n".join((p.extract_text() or "") for p in reader.pages)

    @staticmethod
    def _read_xlsx(path: str) -> str:
        from openpyxl import load_workbook  # type: ignore

        wb = load_workbook(path, read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"【工作表：{ws.title}】")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    # ---------------------------------------------------------------- 检索
    def search(self, query: str, top_k: int = TOP_K) -> list[dict[str, Any]]:
        q = (query or "").strip()
        if not q:
            return []
        # 中文没有空格，整体当成一个子串去匹配；若含空格则按词拆分分别计分
        terms = [t for t in re.split(r"\s+", q) if t]

        with self._conn() as c:
            rows = c.execute(
                "SELECT rowid, doc_id, content FROM chunks_fts"
            ).fetchall()
        scored: list[tuple[int, dict[str, Any]]] = []
        for r in rows:
            content: str = r["content"]
            score = 0
            for t in terms:
                score += content.count(t)
            if score > 0:
                scored.append((score, r))
        scored.sort(key=lambda x: -x[0])

        results: list[dict[str, Any]] = []
        seen_docs: set[int] = set()
        for score, r in scored:
            if r["doc_id"] in seen_docs:
                continue
            seen_docs.add(r["doc_id"])
            doc = self._doc_name(r["doc_id"])
            results.append({"doc": doc, "content": r["content"]})
            if len(results) >= top_k:
                break
        return results

    def _doc_name(self, doc_id: int) -> str:
        with self._conn() as c:
            row = c.execute("SELECT name FROM docs WHERE id=?", (doc_id,)).fetchone()
        return row["name"] if row else f"文档{doc_id}"

    # ---------------------------------------------------------------- 管理
    def list_documents(self) -> list[dict[str, Any]]:
        with self._conn() as c:
            docs = c.execute(
                "SELECT d.id, d.name, d.kind, d.size,"
                " (SELECT COUNT(*) FROM chunks_fts cf WHERE cf.doc_id=d.id) AS chunks"
                " FROM docs d ORDER BY d.created_at DESC"
            ).fetchall()
        return [dict(d) for d in docs]

    def delete_document(self, doc_id: int) -> None:
        with self._conn() as c:
            c.execute("DELETE FROM chunks_fts WHERE doc_id=?", (doc_id,))
            c.execute("DELETE FROM docs WHERE id=?", (doc_id,))

    def clear_all(self) -> None:
        with self._conn() as c:
            c.execute("DELETE FROM chunks_fts")
            c.execute("DELETE FROM docs")

    def get_document_text(self, doc_id: int) -> str:
        """取出某文档的全部文本（按块顺序拼接），供摘要 / 速览使用。"""
        with self._conn() as c:
            rows = c.execute(
                "SELECT content FROM chunks_fts WHERE doc_id=? ORDER BY idx", (doc_id,)
            ).fetchall()
        return "\n\n".join(r["content"] for r in rows)
